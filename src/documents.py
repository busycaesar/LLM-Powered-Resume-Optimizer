from docx import Document

# File Paths.
job_description_file_path = "../documents/job_description.txt"
resume_file_path = "../documents/resume.docx"
projects_file_path = "../documents/projects.docx"

# Get job description content.
def get_job_description():
    with open(job_description_file_path,"r") as file:
        job_description = file.read()
    
    return job_description

# Get resume content.
def get_resume():
    # Takes the resume
    doc = Document(resume_file_path)
    # Extract text from each paragraph and join it into one string
    resume = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    return resume

# Get list of projects.
def get_projects():
    # Takes the document
    doc = Document(projects_file_path)
    # Extract text from each paragraph and join it into one string
    projects = "\n".join([paragraph.text for paragraph in doc.paragraphs])

    return projects